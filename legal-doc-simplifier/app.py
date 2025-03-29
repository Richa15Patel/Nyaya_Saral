from flask import Flask, request, jsonify, render_template, send_file
from huggingface_hub import InferenceClient
import os
import io
import docx
import fitz  # PyMuPDF for PDFs
import pytesseract  # OCR for scanned documents
from PIL import Image
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
import torch
import requests
import json
from IndicTransToolkit import IndicProcessor
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
import threading

app = Flask(__name__)

# Load environment variables
load_dotenv()
HF_API_KEY = os.getenv("HF_API_KEY")
DUCKDUCKGO_API = "https://api.duckduckgo.com/"

# Initialize Hugging Face Inference Client
MODEL_ID = "mistralai/Mistral-7B-Instruct-v0.3"
client = InferenceClient(model=MODEL_ID, token=HF_API_KEY)

# Set upload folder
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

CHUNK_SIZE = 1000  # Chunk size for processing long texts

# Setup IndicTrans2 model and toolkit
model_name = "ai4bharat/indictrans2-en-indic-1B"  # English to Indic model
tokenizer = None
model = None
indic_processor = None
translation_model_lock = threading.Lock()
translation_model_loaded = False

def load_translation_model():
    global tokenizer, model, indic_processor, translation_model_loaded
    with translation_model_lock:
        if not translation_model_loaded:
            DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
            tokenizer = AutoTokenizer.from_pretrained(model_name, trust_remote_code=True)
            model = AutoModelForSeq2SeqLM.from_pretrained(model_name, trust_remote_code=True)
            model.to(DEVICE)
            indic_processor = IndicProcessor(inference=True)
            translation_model_loaded = True

# Start loading the translation model in a background thread
threading.Thread(target=load_translation_model).start()

# Supported Indian languages
SUPPORTED_LANGUAGES = {
    "hindi": "hin_Deva",
    "tamil": "tam_Taml",
    "telugu": "tel_Telu",
    "bengali": "ben_Beng",
    "marathi": "mar_Deva",
    "kannada": "kan_Knda",
    "malayalam": "mal_Mlym",
    "punjabi": "pan_Guru",
    "gujarati": "guj_Gujr",
    "odia": "ory_Orya",
    "assamese": "asm_Beng",
    "bodo": "brx_Deva",
    "dogri": "doi_Deva",
    "konkani": "gom_Deva",
    "kashmiri": "kas_Deva",
    "maithili": "mai_Deva",
    "manipuri": "mni_Beng",
    "meitei": "mni_Mtei",
    "nepali": "npi_Deva",
    "sanskrit": "san_Deva",
    "santali": "sat_Olck",
    "sindhi": "snd_Deva",
    "urdu": "urd_Arab"
}

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/simplify', methods=['POST'])
def simplify_text():
    """Handles text input and file uploads for simplification and translation."""
    
    text = ""
    target_lang = request.form.get("target_lang", "").lower()  
    tgt_lang_code = SUPPORTED_LANGUAGES.get(target_lang)
    
    # Handle file upload
    if "file" in request.files:
        uploaded_file = request.files["file"]
        
        if uploaded_file.filename == "":
            return jsonify({"error": "No file selected"}), 400

        # Save the file
        filename = secure_filename(uploaded_file.filename)
        file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        uploaded_file.save(file_path)

        # Extract text based on file type
        text = extract_text_from_file(file_path, filename)
    
    # Handle direct text input
    elif request.is_json:
        data = request.get_json()
        text = data.get("text", "")
    else:
        # Handle form data
        text = request.form.get("text", "")
    
    if not text:
        return jsonify({"error": "No text provided"}), 400

    # Process text in chunks
    simplified_text = process_text_in_chunks(text)

    # Translate if a target language is chosen
    if tgt_lang_code:
        translated_text = translate_text(simplified_text, tgt_lang_code)
    else:
        translated_text = simplified_text

    # Save the final output as a document
    output_path = save_text_to_docx(translated_text)

    return send_file(output_path, as_attachment=True, download_name="simplified_document.docx")

def extract_text_from_file(file_path, filename):
    """Extracts text from a DOCX or PDF file."""
    if filename.lower().endswith(".docx"):
        return extract_text_from_docx(file_path)
    elif filename.lower().endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError("Unsupported file format")

def extract_text_from_docx(file_path):
    """Extracts text from a .docx file."""
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_pdf(file_path):
    """Extracts text from a PDF, including OCR for scanned documents."""
    text = ""
    doc = fitz.open(file_path)

    for page in doc:
        page_text = page.get_text()
        text += page_text + "\n" if page_text.strip() else perform_ocr_on_pdf_page(page, doc)

    return text.strip()

def perform_ocr_on_pdf_page(page, doc):
    """Performs OCR on scanned PDF pages."""
    text = ""
    img_list = page.get_images(full=True)
    for img in img_list:
        xref = img[0]
        base_image = doc.extract_image(xref)
        img_bytes = base_image["image"]
        img_pil = Image.open(io.BytesIO(img_bytes))
        text += pytesseract.image_to_string(img_pil) + "\n"
    return text

def save_text_to_docx(text):
    """Saves text to a .docx file."""
    output_path = os.path.join(app.config["UPLOAD_FOLDER"], "simplified_document.docx")
    doc = docx.Document()
    
    # Split by paragraphs to preserve formatting
    paragraphs = text.split('\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para)
    
    doc.save(output_path)
    return output_path

def process_text_in_chunks(text):
    """Splits long text into manageable chunks and processes each chunk separately."""
    chunks = chunk_text(text, CHUNK_SIZE)
    return "\n".join([simplify_with_llm(chunk) for chunk in chunks])

def chunk_text(text, chunk_size):
    """Splits text into smaller chunks to avoid exceeding model limits.
    Tries to split at paragraph boundaries when possible."""
    if len(text) <= chunk_size:
        return [text]
    
    # Split by paragraphs first
    paragraphs = text.split("\n")
    chunks = []
    current_chunk = ""
    
    for para in paragraphs:
        # If adding this paragraph would exceed chunk size, save current chunk and start a new one
        if len(current_chunk) + len(para) + 1 > chunk_size:
            if current_chunk:
                chunks.append(current_chunk)
            
            # If a single paragraph is longer than chunk_size, split it
            if len(para) > chunk_size:
                # Split the paragraph into segments
                para_chunks = [para[i:i + chunk_size] for i in range(0, len(para), chunk_size)]
                chunks.extend(para_chunks[:-1])  # Add all but the last segment
                current_chunk = para_chunks[-1]  # Start new chunk with last segment
            else:
                current_chunk = para
        else:
            if current_chunk:
                current_chunk += "\n" + para
            else:
                current_chunk = para
    
    # Add the last chunk if it has content
    if current_chunk:
        chunks.append(current_chunk)
        
    return chunks

def simplify_with_llm(text):
    """Uses Hugging Face's Inference API to simplify legal text."""
    prompt = create_simplification_prompt(text)
    response = client.text_generation(prompt, max_new_tokens=4096, temperature=0.2)
    return response.strip()

def create_simplification_prompt(text):
    """Creates a structured prompt for Mistral to simplify legal text."""
    prompt = f"""<s>[INST] You are a legal text simplification expert. Your task is to convert complex legal language into plain, simple language that an average person with no legal background can understand.

Here are the rules for simplification:
1. Replace legal jargon with everyday language.
2. Break down long sentences into shorter ones.
3. Use active voice instead of passive voice.
4. Maintain the complete meaning and all important details.
5. Do not add any interpretations or new information.
6. Make sure all legal points and rights are preserved.

Please simplify the following legal text:

{text}

Provide only the simplified text in your response, without any additional comments or explanations. [/INST]
"""
    return prompt


def translate_text(text, tgt_lang_code):
    """Translates text to the target Indian language using the ai4bharat model.
    Returns the original text if translation fails."""
    if not tgt_lang_code:
        return text
    
    # Make sure the model is loaded - load on demand if not already loaded
    if not translation_model_loaded:
        load_translation_model()
    
    # Use a try-except block to ensure we return something even if translation fails
    try:
        DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
        
        # Split into paragraphs to handle large text
        paragraphs = text.split("\n")
        translated_paragraphs = []
        
        for paragraph in paragraphs:
            if not paragraph.strip():
                translated_paragraphs.append("")
                continue
                
            # Preprocess
            batch = indic_processor.preprocess_batch([paragraph], src_lang="eng_Latn", tgt_lang=tgt_lang_code)
            
            # Tokenize
            inputs = tokenizer(
                batch,
                truncation=True,
                padding="longest",
                return_tensors="pt",
                return_attention_mask=True,
            ).to(DEVICE)
            
            # Generate Translation
            with torch.no_grad():
                generated_tokens = model.generate(
                    **inputs,
                    use_cache=True,
                    min_length=0,
                    max_length=256,
                    num_beams=5,
                    num_return_sequences=1,
                )
            
            # Decode Translation
            with tokenizer.as_target_tokenizer():
                translated_text = tokenizer.batch_decode(
                    generated_tokens.detach().cpu().tolist(),
                    skip_special_tokens=True,
                    clean_up_tokenization_spaces=True,
                )[0]
            
            # Postprocess
            translated_paragraph = indic_processor.postprocess_batch([translated_text], lang=tgt_lang_code)[0]
            translated_paragraphs.append(translated_paragraph)
        
        return "\n".join(translated_paragraphs)
    
    except Exception as e:
        # Log the error but return the original text
        print(f"Translation error: {str(e)}")
        return text + "\n\n[Translation failed. Showing simplified English text.]"

@app.route('/api/languages', methods=['GET'])
def get_languages():
    """Returns the list of supported languages."""
    return jsonify({name: name.capitalize() for name in SUPPORTED_LANGUAGES.keys()})

if __name__ == '__main__':
    app.run(debug=True)